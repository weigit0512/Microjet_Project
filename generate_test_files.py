import os
import pandas as pd
import random
from datetime import datetime, timedelta
from fpdf import FPDF
from docx import Document

# 建立輸出資料夾
OUTPUT_DIR = "addwii_test_suite"
os.makedirs(OUTPUT_DIR, exist_ok=True)

print("開始生成企業級實測檔案...")

# ==========================================
# 場景 C: 真實 PDF 提案生成範例 (B2B_Proposal.pdf)
# ==========================================
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'Microjet B2B Partnership Proposal - MJ3200', 0, 1, 'C')
        self.ln(5)

pdf = PDF()
pdf.add_page()
pdf.set_font('Arial', '', 12)
sections = [
    "1. Executive Summary: Enhancing printing efficiency for specific regions.",
    "2. Collaboration History: 450 units purchased in the last 12 months.",
    "3. Market Analysis: 15% growth potential in the retail sector.",
    "4. New Product Recommendation: MJ-3200 Advanced Inkjet.",
    "5. Procurement Plan: Volume discount applied (Tier 2).",
    "6. Marketing Support: Co-branded POSM provided.",
    "7. Mutual Commitments: 99% SLA guarantee.",
    "8. Appendix: Detailed specs and terms."
]
for sec in sections:
    pdf.cell(0, 10, sec, 0, 1)
pdf.output(os.path.join(OUTPUT_DIR, "Scenario_C_B2B_Proposal.pdf"))
print("✅ 已生成: PDF 提案範本")

# ==========================================
# 場景 D: 客戶回饋 CSV 檔 (Customer_Feedback.csv)
# ==========================================
dates = [(datetime.now() - timedelta(days=i)).strftime("%Y-%m-%d") for i in range(100)]
platforms = ["PChome", "Momo", "Shopee", "Official_Site"]
feedbacks = [
    ("墨水匣很容易乾掉，太誇張了", 1), ("Wi-Fi 一直斷線，連不上", 2), 
    ("出貨等了快一個禮拜", 2), ("列印速度真的很快，滿意！", 5), 
    ("保固服務很棒，客服很親切", 5), ("相容性很高，Mac也能順跑", 4),
    ("更新 v2.15 後一直卡紙 (E-044)", 1), ("機器不錯但耗材太貴", 3)
]

csv_data = []
for _ in range(500): # 生成 500 筆測試資料
    fb = random.choice(feedbacks)
    csv_data.append({
        "Date": random.choice(dates),
        "Platform": random.choice(platforms),
        "Rating": fb[1],
        "Comment": fb[0]
    })
pd.DataFrame(csv_data).to_csv(os.path.join(OUTPUT_DIR, "Scenario_D_Customer_Feedback.csv"), index=False, encoding='utf-8-sig')
print("✅ 已生成: 500 筆客戶回饋 CSV 檔")

# ==========================================
# 場景 E: 混合測試檔組 (Word, Excel, Log 埋設 PII)
# ==========================================
pii_dir = os.path.join(OUTPUT_DIR, "Scenario_E_Mixed_PII_Suite")
os.makedirs(pii_dir, exist_ok=True)

# 1. 生成 Word (埋設身分證與手機)
doc = Document()
doc.add_heading('員工健檢名單 (內部機密)', 0)
doc.add_paragraph('王大明, ID: A123456789, Phone: 0912-345-678')
doc.add_paragraph('陳小美, ID: B234567890, Phone: 0987-654-321')
# 模擬塞入 40 筆身分證...
doc.save(os.path.join(pii_dir, "Employee_Health_List.docx"))

# 2. 生成 Excel (埋設信用卡與地址)
excel_data = {
    "Customer": ["張先生", "林小姐", "李總"],
    "Credit_Card": ["1234-5678-1234-5678", "4321-8765-4321-8765", "5555-6666-7777-8888"],
    "Address": ["台北市信義區松智路1號", "新北市板橋區縣民大道2段", "台中市西屯區台灣大道3段"]
}
pd.DataFrame(excel_data).to_excel(os.path.join(pii_dir, "VIP_Billing_Info.xlsx"), index=False)

# 3. 生成 Access Log (埋設存取異常)
with open(os.path.join(pii_dir, "server_access.log"), "w", encoding="utf-8") as f:
    f.write("2026-04-15 09:12:33 INFO User_A logged in\n")
    f.write("2026-04-15 14:22:11 INFO User_B downloaded 15 records\n")
    # 埋設異常紀錄
    f.write("2026-04-15 23:47:05 WARN User_X downloaded 1284 records from /api/customers [OUT_OF_OFFICE_HOURS]\n")

print("✅ 已生成: 包含 PII 與異常行為的混合測試檔組")
print(f"🎉 所有測試檔案皆已保存在 '{OUTPUT_DIR}' 資料夾中！")